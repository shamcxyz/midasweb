import os
import tempfile
import zipfile
from datetime import datetime
from typing import List, Dict, Union
from fastapi import FastAPI, HTTPException, Form, File, UploadFile, status
from dotenv import load_dotenv
from openai import OpenAI
import base64
import logging
import boto3
from botocore.exceptions import NoCredentialsError, ClientError
from docx import Document
import PyPDF2
import re  # Import the re module for filename sanitization

# Initialize Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load Environment Variables
load_dotenv()

# OpenAI Configuration
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# AWS S3 Configuration
AWS_ACCESS_KEY_ID = os.getenv("AWS_ACCESS_KEY_ID")
AWS_SECRET_ACCESS_KEY = os.getenv("AWS_SECRET_ACCESS_KEY")
AWS_S3_BUCKET_NAME = os.getenv("AWS_S3_BUCKET_NAME", "midasbucket")
AWS_REGION = os.getenv("AWS_REGION", "us-east-2")

# Initialize OpenAI client
client = OpenAI(api_key=OPENAI_API_KEY)

# Initialize S3 client using Main Account Credentials
s3_client = boto3.client(
    's3',
    region_name=AWS_REGION,
    aws_access_key_id=AWS_ACCESS_KEY_ID,
    aws_secret_access_key=AWS_SECRET_ACCESS_KEY
)

app = FastAPI(title="Reimbursement Request System")

class MultiFileProcessor:
    ALLOWED_EXTENSIONS = {'.docx', '.pdf', '.jpg', '.jpeg', '.png', '.zip'}
    IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png'}
    
    @staticmethod
    def get_file_extension(filename: str) -> str:
        return os.path.splitext(filename)[1].lower()
    
    @staticmethod
    def encode_image(image_path: str) -> str:
        try:
            with open(image_path, "rb") as image_file:
                return base64.b64encode(image_file.read()).decode('utf-8')
        except Exception as e:
            logger.error(f"Error encoding image: {str(e)}")
            raise ValueError(f"Error encoding image: {str(e)}")

    @staticmethod
    def extract_text_from_pdf(pdf_path: str) -> str:
        try:
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                content = []
                
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        content.append(text)
                        
                text = "\n\n=== Page Break ===\n\n".join(content)
                
                lines = text.split('\n')
                formatted_lines = []
                for line in lines:
                    if '\t' in line or '    ' in line:
                        items = [item.strip() for item in line.split('\t') if item.strip()]
                        if not items:
                            items = [item.strip() for item in line.split('    ') if item.strip()]
                        if items:
                            formatted_lines.append('\t'.join(items))
                    else:
                        formatted_lines.append(line)
                        
                return '\n'.join(formatted_lines)
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise ValueError(f"Error extracting text from PDF: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        try:
            with open(file_path, 'rb') as file:
                doc = Document(file)
                content = []
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        content.append(para.text.strip())
                
                for table in doc.tables:
                    table_content = []
                    for row in table.rows:
                        row_content = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_content:
                            table_content.append('\t'.join(row_content))
                    if table_content:
                        content.append('\n'.join(table_content))
                
                return '\n\n'.join(content)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise ValueError(f"Error extracting text from DOCX: {str(e)}")

    @staticmethod
    async def save_upload_file(upload_file: UploadFile) -> str:
        try:
            # Use the original filename for the temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(upload_file.filename)[1]) as tmp_file:
                content = await upload_file.read()
                tmp_file.write(content)
                return tmp_file.name
        except Exception as e:
            logger.error(f"Error saving uploaded file: {str(e)}")
            raise ValueError(f"Error saving uploaded file: {str(e)}")

    @staticmethod
    def process_zip(zip_path: str) -> List[Dict[str, Union[str, bool]]]:
        processed_files = []
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(tmp_dir)
                    
                    for root, _, files in os.walk(tmp_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            ext = MultiFileProcessor.get_file_extension(file)
                            
                            if ext in MultiFileProcessor.ALLOWED_EXTENSIONS and ext != '.zip':
                                try:
                                    processed_content = MultiFileProcessor.process_single_file(file_path)
                                    processed_files.append(processed_content)
                                except Exception as e:
                                    logger.error(f"Error processing file {file} in zip: {str(e)}")
                                    continue
        except Exception as e:
            logger.error(f"Error processing zip file: {str(e)}")
            raise ValueError(f"Error processing zip file: {str(e)}")
        
        return processed_files

    @staticmethod
    def process_single_file(file_path: str) -> Dict[str, Union[str, bool]]:
        ext = MultiFileProcessor.get_file_extension(file_path)
        
        try:
            if ext == '.docx':
                return {
                    'content': MultiFileProcessor.extract_text_from_docx(file_path),
                    'is_image': False
                }
            elif ext == '.pdf':
                return {
                    'content': MultiFileProcessor.extract_text_from_pdf(file_path),
                    'is_image': False
                }
            elif ext in MultiFileProcessor.IMAGE_EXTENSIONS:
                return {
                    'content': MultiFileProcessor.encode_image(file_path),
                    'is_image': True
                }
            else:
                raise ValueError(f"Unsupported file type: {ext}")
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            raise ValueError(f"Error processing file {file_path}: {str(e)}")

async def analyze_with_gpt4o(content: str, is_image: bool = False) -> Dict[str, str]:
    current_date = datetime.now().strftime("%Y-%m-%d")
    
    try:
        if is_image:
            messages = [
                {
                    "role": "system",
                    "content": (
                        "You are an assistant helping with reimbursement requests. "
                        "Analyze the receipt image and decide whether to 'Approve' or 'Reject' the request. "
                        "Please respond in the following format:\n\n"
                        "Decision: [Approve/Reject]\n"
                        "Feedback: [Your explanation here]\n\n"
                        "Please avoid using ambiguous language that might put the decision in doubt."
                    )
                },
                {
                    "role": "user",
                    "content": f"Today is {current_date}. Please analyze this receipt image in base64 format:\n\n{content}"
                }
            ]
        else:
            messages = [
                {
                    "role": "system",
                    "content": (
                        "You are an assistant helping with reimbursement requests. "
                        "Analyze the document and decide whether to 'Approve' or 'Reject' the request. "
                        "Please respond in the following format:\n\n"
                        "Decision: [Approve/Reject]\n"
                        "Feedback: [Your explanation here]\n\n"
                        "Please avoid using ambiguous language that might put the decision in doubt."
                    )
                },
                {
                    "role": "user",
                    "content": f"Today is {current_date}. Here is the document content:\n\n{content}"
                }
            ]
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages
        )
        
        analysis = response.choices[0].message.content

        # Parse the decision and feedback from the analysis
        lines = analysis.strip().split('\n')
        decision = None
        feedback_lines = []
        for line in lines:
            if line.lower().startswith("decision:"):
                decision = line[len("Decision:"):].strip()
            elif line.lower().startswith("feedback:"):
                feedback_lines.append(line[len("Feedback:"):].strip())
            else:
                feedback_lines.append(line.strip())
        feedback = '\n'.join(feedback_lines)

        return {'decision': decision, 'feedback': feedback}
                
    except Exception as e:
        logger.error(f"Error analyzing content: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error analyzing content: {str(e)}"
        )

async def verify_decision_feedback(decision: str, feedback: str) -> bool:
    try:
        messages = [
            {
                "role": "system",
                "content": (
                    "You are an assistant that verifies whether the provided decision aligns with the feedback. "
                    "Respond with 'Yes' if they match or 'No' if they don't."
                )
            },
            {
                "role": "user",
                "content": (
                    f"Decision: {decision}\n"
                    f"Feedback: {feedback}\n\n"
                    "Does the decision match the feedback?"
                )
            }
        ]

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages
        )

        verification = response.choices[0].message.content.strip().lower()
        if 'yes' in verification:
            return True
        else:
            return False

    except Exception as e:
        logger.error(f"Error verifying decision and feedback: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error verifying decision and feedback: {str(e)}"
        )

def sanitize_filename(filename):
    # Remove any path separators
    filename = os.path.basename(filename)
    # Remove any potentially harmful characters
    filename = re.sub(r'[^\w\.-]', '_', filename)
    return filename

def upload_to_s3(file_path: str, original_filename: str, decision: str, admin_email: str) -> str:
    try:
        target_repo = admin_email.replace('@', '').replace('.', '')

        name, ext = os.path.splitext(original_filename)
        current_datetime = datetime.now().strftime("%Y%m%d_%H%M%S")
        new_filename = f"{name}_{current_datetime}{ext}"
        object_name = f"Reimbursement/{target_repo}/{decision.upper()}/{new_filename}"

        s3_client.upload_file(file_path, AWS_S3_BUCKET_NAME, object_name)
        s3_url = f"https://{AWS_S3_BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{object_name}"
        logger.info(f"Uploaded {file_path} to {s3_url}")
        return s3_url
    except FileNotFoundError:
        logger.error(f"The file {file_path} was not found.")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"The file {file_path} was not found."
        )
    except NoCredentialsError:
        logger.error("AWS credentials not available.")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="AWS credentials not available."
        )
    except ClientError as e:
        logger.error(f"Client error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Error uploading file to S3."
        )

@app.post("/request_reimbursement")
async def request_reimbursement(
    role: str = Form(...),
    name: str = Form(...),
    email: str = Form(...),
    admin_email: str = Form(...),
    reimbursement_details: str = Form(...),
    files: List[UploadFile] = File(...)
):
    if role.lower() != "user":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Role must be 'user'."
        )
    
    processor = MultiFileProcessor()
    temp_files = []
    all_content = []
    s3_urls = []
    
    try:
        for file in files:
            ext = processor.get_file_extension(file.filename)
            if ext not in processor.ALLOWED_EXTENSIONS:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Unsupported file type: {ext}"
                )
            
            temp_file_path = await processor.save_upload_file(file)
            original_filename = sanitize_filename(file.filename)
            temp_files.append((temp_file_path, original_filename))  # Keep sanitized original filename
            
            if ext == '.zip':
                zip_contents = processor.process_zip(temp_file_path)
                all_content.extend(zip_contents)
            else:
                processed_content = processor.process_single_file(temp_file_path)
                all_content.append(processed_content)
        
        combined_feedback = ""
        decisions = []
        for content_item in all_content:
            analysis_result = await analyze_with_gpt4o(
                content_item['content'],
                is_image=content_item.get('is_image', False)
            )
            decision = analysis_result['decision']
            feedback = analysis_result['feedback']

            # Verify decision and feedback
            is_verified = await verify_decision_feedback(decision, feedback)
            if not is_verified:
                # Adjust the decision based on verification
                decision = 'Rejected'

            combined_feedback += f"{feedback}\n\n"
            decisions.append(decision.capitalize())
        
        # Determine the overall decision
        if all(dec.lower() == 'approved' for dec in decisions):
            final_decision = 'Approved'
        else:
            final_decision = 'Rejected'
        
        # Upload files to S3 with appropriate naming
        for temp_file, original_filename in temp_files:
            s3_url = upload_to_s3(temp_file, original_filename, final_decision, admin_email)
            s3_urls.append(s3_url)
        
        return {
            "status": final_decision,
            "feedback": combined_feedback.strip(),
            "processed_files": len(all_content),
            "uploaded_files": s3_urls
        }
            
    except HTTPException as he:
        logger.error(f"HTTPException: {he.detail}")
        raise he
    except Exception as e:
        logger.error(f"Error in request_reimbursement: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error processing request: {str(e)}"
        )
    finally:
        for temp_file, _ in temp_files:
            if os.path.exists(temp_file):
                os.remove(temp_file)
